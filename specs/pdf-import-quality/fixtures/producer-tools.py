#!/usr/bin/env python3
"""Prepare and sanitize customer-neutral desktop producer fixtures."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject


HERE = Path(__file__).resolve().parent
SOURCE_PATH = HERE / "producer-source.json"


def configure_run(run: object) -> None:
    run.font.name = "Arial Unicode MS"
    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    font_name = "MS Mincho" if "港" in run.text else "Arial Unicode MS"
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), font_name)
    if "港" in run.text:
        language = properties.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            properties.append(language)
        language.set(qn("w:eastAsia"), "ja-JP")
    if "مرحبا" in run.text:
        language = properties.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            properties.append(language)
        language.set(qn("w:bidi"), "ar-SA")


def make_producer_docx(path: Path) -> None:
    source = json.loads(SOURCE_PATH.read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    document.core_properties.title = source["title"]
    document.core_properties.author = "AtlCLI neutral fixture generator"
    document.core_properties.subject = "Synthetic PDF import quality evidence"
    document.add_heading(source["title"], level=1)
    for paragraph_runs in source["paragraphs"]:
        paragraph = document.add_paragraph()
        for entry in paragraph_runs:
            run = paragraph.add_run(entry["text"])
            run.bold = bool(entry.get("bold", False))
            run.font.size = Pt(11)
        if any("مرحبا" in entry["text"] for entry in paragraph_runs):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for item in source["listItems"]:
        document.add_paragraph(item, style="List Number")
    table = document.add_table(rows=len(source["table"]), cols=2)
    table.style = "Table Grid"
    for row_index, values in enumerate(source["table"]):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = value
            if row_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            configure_run(run)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        configure_run(run)
    document.save(path)


def sanitize_export(input_path: Path, output_path: Path, producer: str) -> None:
    reader = PdfReader(input_path)
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    writer.metadata = None
    writer._root_object.pop(NameObject("/Metadata"), None)
    writer.add_metadata(
        {
            "/Title": "Neutral Harbor Field Notes",
            "/Author": "AtlCLI neutral fixture generator",
            "/Subject": "Synthetic PDF import quality evidence",
            "/Creator": producer,
            "/Producer": producer,
        }
    )
    writer.compress_identical_objects(remove_duplicates=True, remove_unreferenced=True)
    with output_path.open("wb") as output:
        writer.write(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest="command", required=True)
    producer = subparsers.add_parser("producer-source")
    producer.add_argument("--output", type=Path, required=True)
    sanitize = subparsers.add_parser("sanitize-export")
    sanitize.add_argument("--input", type=Path, required=True)
    sanitize.add_argument("--output", type=Path, required=True)
    sanitize.add_argument("--producer", required=True)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.command == "producer-source":
        make_producer_docx(args.output)
        print(f"created neutral producer source: {args.output}")
    elif args.command == "sanitize-export":
        sanitize_export(args.input, args.output, args.producer)
        print(f"sanitized neutral producer export: {args.output}")


if __name__ == "__main__":
    main()
