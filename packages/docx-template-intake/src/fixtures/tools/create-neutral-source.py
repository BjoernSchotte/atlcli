#!/usr/bin/env python3
"""Create the synthetic source used for the T0 Word/LibreOffice fixtures.

The output deliberately contains only marker text listed in ../README.md. It is
saved once by Microsoft Word and once by LibreOffice before those two binaries
are committed. The script itself does not call either office application.
"""

from __future__ import annotations

import argparse
import os
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


MARKERS = {
    "title": "NEUTRAL_FIXTURE_TITLE",
    "subtitle": "NEUTRAL_FIXTURE_SUBTITLE",
    "heading_1": "NEUTRAL_HEADING_ONE",
    "heading_2": "NEUTRAL_HEADING_TWO",
    "heading_3": "NEUTRAL_HEADING_THREE",
    "body_alpha": "NEUTRAL_BODY_ALPHA",
    "body_beta": "NEUTRAL_BODY_BETA",
    "code": "NEUTRAL_CODE_SAMPLE",
    "caption": "NEUTRAL_CAPTION",
    "table_header_a": "NEUTRAL_TABLE_HEADER_A",
    "table_header_b": "NEUTRAL_TABLE_HEADER_B",
    "table_cell_a": "NEUTRAL_TABLE_CELL_A",
    "table_cell_b": "NEUTRAL_TABLE_CELL_B",
    "header_default": "NEUTRAL_HEADER_DEFAULT",
    "header_first": "NEUTRAL_HEADER_FIRST",
    "header_even": "NEUTRAL_HEADER_EVEN",
    "footer": "NEUTRAL_FOOTER",
    "section_two": "NEUTRAL_SECTION_TWO",
}

INK = RGBColor(0x18, 0x22, 0x33)
ACCENT = RGBColor(0x4B, 0x57, 0xA3)
MUTED = RGBColor(0x5D, 0x68, 0x7A)


def set_run_font(run, name: str, size: float, color: RGBColor, *, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_page_border(section) -> None:
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(existing)
    page_borders = OxmlElement("w:pgBorders")
    page_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), "D9DDEA")
        page_borders.append(border)
    sect_pr.append(page_borders)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Aptos")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (18, 16, 8),
        "Heading 2": (14, 12, 6),
        "Heading 3": (12, 9, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Aptos Display")
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code = styles.add_style("Neutral Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Courier New"
    code._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Courier New")
    code._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(9)
    code.font.color.rgb = INK
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)

    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED


def make_images(directory: Path) -> tuple[Path, Path, bytes]:
    png_path = directory / "neutral-mark.png"
    png = Image.new("RGBA", (640, 180), (75, 87, 163, 255))
    draw = ImageDraw.Draw(png)
    draw.rectangle((36, 36, 144, 144), fill=(247, 248, 252, 255))
    draw.rectangle((176, 54, 584, 126), fill=(24, 34, 51, 255))
    png.save(png_path, format="PNG", optimize=True)

    jpeg_path = directory / "neutral-photo.jpg"
    jpeg = Image.new("RGB", (720, 420), (241, 233, 218))
    draw = ImageDraw.Draw(jpeg)
    draw.rectangle((50, 50, 670, 370), fill=(217, 221, 234), outline=(75, 87, 163), width=8)
    draw.ellipse((250, 100, 470, 320), fill=(75, 87, 163), outline=(24, 34, 51), width=8)
    jpeg.save(jpeg_path, format="JPEG", quality=88, optimize=True)

    svg = (
        '<svg xmlns="http://www.w3.org/2000/svg" width="640" height="180" '
        'viewBox="0 0 640 180">'
        '<rect width="640" height="180" fill="#4B57A3"/>'
        '<path d="M36 36H144V144H36Z" fill="#F7F8FC"/>'
        '<path d="M176 54H584V126H176Z" fill="#182233"/>'
        "</svg>"
    ).encode("utf-8")
    return png_path, jpeg_path, svg


def add_header_picture(header, png_path: Path) -> None:
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    picture_run = paragraph.add_run()
    picture_run.add_picture(str(png_path), width=Inches(1.45))
    label_run = paragraph.add_run(f"  {MARKERS['header_default']}")
    set_run_font(label_run, "Aptos", 8.5, MUTED, bold=True)


def configure_headers_and_footers(doc: Document, png_path: Path) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))

    first = doc.sections[0]
    first.different_first_page_header_footer = True
    add_header_picture(first.header, png_path)

    first_header = first.first_page_header.paragraphs[0]
    first_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(first_header.add_run(MARKERS["header_first"]), "Aptos", 8.5, MUTED, bold=True)

    even_header = first.even_page_header.paragraphs[0]
    even_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(even_header.add_run(MARKERS["header_even"]), "Aptos", 8.5, MUTED, bold=True)

    footer = first.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run(MARKERS["footer"]), "Aptos", 8, MUTED)


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_page_border(section)


def create_document(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="atlcli-neutral-docx-") as tmp:
        tmp_dir = Path(tmp)
        png_path, jpeg_path, svg_bytes = make_images(tmp_dir)

        doc = Document()
        doc.core_properties.author = "ATLCLI Synthetic Fixture"
        doc.core_properties.last_modified_by = "ATLCLI Synthetic Fixture"
        doc.core_properties.title = "Neutral DOCX Template Fixture"
        doc.core_properties.subject = "ATLCLI_SYNTHETIC_FIXTURE"
        doc.core_properties.created = datetime(2020, 1, 1, tzinfo=timezone.utc)
        doc.core_properties.modified = datetime(2020, 1, 1, tzinfo=timezone.utc)

        configure_styles(doc)
        configure_page(doc.sections[0])
        configure_headers_and_footers(doc, png_path)

        title = doc.add_paragraph()
        title.paragraph_format.space_before = Pt(22)
        title.paragraph_format.space_after = Pt(4)
        set_run_font(title.add_run(MARKERS["title"]), "Aptos Display", 28, INK, bold=True)

        subtitle = doc.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(18)
        set_run_font(subtitle.add_run(MARKERS["subtitle"]), "Aptos", 13, MUTED)

        doc.add_heading(MARKERS["heading_1"], level=1)
        doc.add_paragraph(MARKERS["body_alpha"])
        direct = doc.add_paragraph()
        direct.paragraph_format.space_after = Pt(10)
        set_run_font(direct.add_run(MARKERS["body_beta"]), "Aptos", 11.5, ACCENT, bold=True)

        doc.add_heading(MARKERS["heading_2"], level=2)
        doc.add_paragraph(MARKERS["code"], style="Neutral Code")
        doc.add_heading(MARKERS["heading_3"], level=3)

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        values = (
            (MARKERS["table_header_a"], MARKERS["table_header_b"]),
            (MARKERS["table_cell_a"], MARKERS["table_cell_b"]),
        )
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.width = Inches(3.15)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = values[row_index][column_index]
                if row_index == 0:
                    set_cell_shading(cell, "E8EAF4")
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

        doc.add_paragraph().add_run().add_picture(str(jpeg_path), width=Inches(3.6))
        doc.add_paragraph(MARKERS["caption"], style="Caption")

        second = doc.add_section(WD_SECTION.NEW_PAGE)
        second.orientation = WD_ORIENT.LANDSCAPE
        second.page_width = Inches(11)
        second.page_height = Inches(8.5)
        second.top_margin = Inches(0.75)
        second.right_margin = Inches(0.75)
        second.bottom_margin = Inches(0.75)
        second.left_margin = Inches(0.75)
        second.header_distance = Inches(0.35)
        second.footer_distance = Inches(0.35)
        add_page_border(second)
        doc.add_heading(MARKERS["section_two"], level=1)
        doc.add_paragraph(MARKERS["body_alpha"])

        doc.save(output)
        patch_package(output, svg_bytes)


def patch_package(path: Path, svg_bytes: bytes) -> None:
    with zipfile.ZipFile(path, "r") as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]

    rewritten: list[tuple[zipfile.ZipInfo, bytes]] = []
    header_patched = False
    relationship_patched = False
    svg_content_type_patched = False

    for info, data in entries:
        if info.filename == "[Content_Types].xml":
            text = data.decode("utf-8")
            if 'Extension="svg"' not in text:
                text = text.replace(
                    "</Types>",
                    '<Default Extension="svg" ContentType="image/svg+xml"/></Types>',
                )
            data = text.encode("utf-8")
            svg_content_type_patched = True
        elif info.filename == "word/document.xml":
            text = data.decode("utf-8")
            if "<w:background" not in text:
                text, count = re.subn(
                    r"(<w:document\b[^>]*>)",
                    r'\1<w:background w:color="F7F8FC"/>',
                    text,
                    count=1,
                )
                if count != 1:
                    raise RuntimeError("failed to add the synthetic document background")
            data = text.encode("utf-8")
        elif info.filename == "word/theme/theme1.xml":
            text = data.decode("utf-8")
            replacements = {
                "accent1": "4B57A3",
                "accent2": "182233",
                "accent3": "D9DDEA",
                "accent4": "A9B0C5",
                "accent5": "6D7895",
                "accent6": "E8EAF4",
            }
            for name, value in replacements.items():
                text, count = re.subn(
                    rf'(<a:{name}>\s*<a:srgbClr val=")[0-9A-Fa-f]{{6}}("/>)',
                    rf"\g<1>{value}\2",
                    text,
                    count=1,
                )
                if count != 1:
                    raise RuntimeError(f"failed to replace theme color {name}")
            data = text.encode("utf-8")
        elif info.filename.startswith("word/header") and info.filename.endswith(".xml") and not header_patched:
            text = data.decode("utf-8")
            if "<a:blip " in text:
                text, count = re.subn(
                    r"(<a:blip\b[^>]*/>)",
                    (
                        r'<a:blip r:embed="rIdSvgFallback"><a:extLst>'
                        r'<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
                        r'<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
                        r'r:embed="rIdSvgFixture"/>'
                        r"</a:ext></a:extLst></a:blip>"
                    ),
                    text,
                    count=1,
                )
                if count != 1:
                    raise RuntimeError("failed to add the SVG extension to the header drawing")
                # Preserve the original raster relationship as the explicit fallback.
                original_match = re.search(r'r:embed="([^"]+)"', data.decode("utf-8"))
                if original_match:
                    text = text.replace("rIdSvgFallback", original_match.group(1), 1)
                text = text.replace(
                    "<a:stretch>",
                    '<a:srcRect l="2500" t="0" r="2500" b="0"/><a:stretch>',
                    1,
                )
                header_patched = True
            data = text.encode("utf-8")
        elif (
            info.filename.startswith("word/_rels/header")
            and info.filename.endswith(".xml.rels")
            and header_patched
            and not relationship_patched
        ):
            text = data.decode("utf-8")
            text = text.replace(
                "</Relationships>",
                (
                    '<Relationship Id="rIdSvgFixture" '
                    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                    'Target="media/neutral-mark.svg"/></Relationships>'
                ),
            )
            data = text.encode("utf-8")
            relationship_patched = True
        rewritten.append((info, data))

    if not (header_patched and relationship_patched and svg_content_type_patched):
        raise RuntimeError("failed to add the SVG-with-raster-fallback fixture")

    svg_info = zipfile.ZipInfo("word/media/neutral-mark.svg")
    svg_info.date_time = (2020, 1, 1, 0, 0, 0)
    svg_info.compress_type = zipfile.ZIP_DEFLATED
    rewritten.append((svg_info, svg_bytes))

    temporary = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(temporary, "w") as archive:
        for info, data in rewritten:
            info.date_time = (2020, 1, 1, 0, 0, 0)
            info.create_system = 0
            archive.writestr(info, data)
    os.replace(temporary, path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    create_document(args.output.resolve())


if __name__ == "__main__":
    main()
