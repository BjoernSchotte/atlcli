"""Convert markdown to Word document elements."""

import base64
import io
import re
from typing import Optional
import markdown
from markdown.extensions.tables import TableExtension
from markdown.extensions.fenced_code import FencedCodeExtension
from bs4 import BeautifulSoup, Tag
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate


# Panel colors for info, warning, note, tip macros
PANEL_STYLES = {
    "info": {"bg": "DEEBFF", "border": "0052CC", "icon": "ℹ️"},
    "warning": {"bg": "FFFAE6", "border": "FF8B00", "icon": "⚠️"},
    "note": {"bg": "EAE6FF", "border": "6554C0", "icon": "📝"},
    "tip": {"bg": "E3FCEF", "border": "00875A", "icon": "💡"},
    "error": {"bg": "FFEBE6", "border": "DE350B", "icon": "❌"},
}

# Status badge colors
STATUS_COLORS = {
    "green": RGBColor(0, 128, 0),
    "yellow": RGBColor(204, 153, 0),
    "red": RGBColor(204, 0, 0),
    "blue": RGBColor(0, 102, 204),
    "grey": RGBColor(128, 128, 128),
    "gray": RGBColor(128, 128, 128),
    "purple": RGBColor(128, 0, 128),
}


def preprocess_markdown(md_text: str) -> str:
    """Pre-process custom markdown syntax before standard markdown conversion.

    Converts:
    - :::type Title\\ncontent\\n::: → <div class="panel panel-type"><div class="panel-title">Title</div>content</div>
    - {color:name}[TEXT]{color} → <span class="status status-name">TEXT</span>
    """
    def render_markdown_fragment(text: str) -> str:
        md = markdown.Markdown(
            extensions=[TableExtension(), FencedCodeExtension(), "nl2br"]
        )
        html = md.convert(text)
        md.reset()
        return html

    def parse_param(params: str, key: str) -> str:
        if not params:
            return ""
        # key="value" or key=value
        match = re.search(rf'{key}\s*=\s*("[^"]*"|\S+)', params)
        if not match:
            return ""
        val = match.group(1)
        return val.strip('"')

    result = md_text

    # Convert TOC macro to placeholder div - always inject Word-native TOC field
    result = re.sub(
        r':::toc\s*\n:::',
        '<div class="toc-macro"></div>',
        result
    )

    # Convert expand macro (with body) to HTML div
    expand_pattern = re.compile(
        r':::expand(?: ([^\n]*))?\n(.*?)\n:::',
        re.DOTALL
    )

    def expand_replacer(match):
        title = (match.group(1) or "Click to expand").strip()
        content = match.group(2).strip()
        html = render_markdown_fragment(content)
        return f'<div class="expand" data-title="{title}">{html}</div>'

    result = expand_pattern.sub(expand_replacer, result)

    # Convert excerpt macro (with body) to HTML div (content only)
    excerpt_pattern = re.compile(
        r':::excerpt(?: ([^\n]*))?\n(.*?)\n:::',
        re.DOTALL
    )

    def excerpt_replacer(match):
        params = (match.group(1) or "").strip()
        content = match.group(2).strip()
        name = parse_param(params, "name")
        hidden = "hidden" in params.split()
        html = render_markdown_fragment(content)
        hidden_attr = ' data-hidden="true"' if hidden else ""
        name_attr = f' data-name="{name}"' if name else ""
        return f'<div class="excerpt"{name_attr}{hidden_attr}>{html}</div>'

    result = excerpt_pattern.sub(excerpt_replacer, result)

    # Convert children macro to placeholder div
    children_pattern = re.compile(
        r':::children(?: ([^\n]*))?\n:::',
        re.DOTALL
    )

    def children_replacer(match):
        params = (match.group(1) or "").strip()
        depth = parse_param(params, "depth")
        depth_attr = f' data-depth="{depth}"' if depth else ""
        return f'<div class="children"{depth_attr}></div>'

    result = children_pattern.sub(children_replacer, result)

    # Convert content-by-label macro to placeholder div
    cbl_pattern = re.compile(
        r':::content-by-label(?: ([^\n]*))?\n:::',
        re.DOTALL
    )

    def cbl_replacer(match):
        params = (match.group(1) or "").strip()
        labels = parse_param(params, "labels")
        spaces = parse_param(params, "spaces")
        max_val = parse_param(params, "max")
        attrs = []
        if labels:
            attrs.append(f'data-labels="{labels}"')
        if spaces:
            attrs.append(f'data-spaces="{spaces}"')
        if max_val:
            attrs.append(f'data-max="{max_val}"')
        attrs_str = " " + " ".join(attrs) if attrs else ""
        return f'<div class="content-by-label"{attrs_str}></div>'

    result = cbl_pattern.sub(cbl_replacer, result)

    # Convert page-properties macro to HTML table
    page_props_pattern = re.compile(
        r':::page-properties(?: ([^\n]*))?\n(.*?)\n:::',
        re.DOTALL
    )

    def page_props_replacer(match):
        params = (match.group(1) or "").strip()
        content = match.group(2).strip()
        html = render_markdown_fragment(content)
        data_id = parse_param(params, "id")
        id_attr = f' data-id="{data_id}"' if data_id else ""
        return f'<div class="page-properties"{id_attr}>{html}</div>'

    result = page_props_pattern.sub(page_props_replacer, result)

    # Convert panel macros: :::type Title\ncontent\n:::
    # Pattern matches :::type optional-title\ncontent\n::: (non-greedy, stopping at first \n:::)
    panel_types = "|".join(PANEL_STYLES.keys())
    panel_pattern = re.compile(
        rf':::({panel_types})(?: ([^\n]*))?\n(.*?)\n:::',
        re.DOTALL
    )

    def panel_replacer(match):
        panel_type = match.group(1).lower()
        title = (match.group(2) or "").strip()
        content = match.group(3).strip()

        if title:
            return f'<div class="panel panel-{panel_type}"><div class="panel-title">{title}</div>\n\n{content}\n\n</div>'
        else:
            return f'<div class="panel panel-{panel_type}">\n\n{content}\n\n</div>'

    result = panel_pattern.sub(panel_replacer, result)

    # Convert status badges: {color:name}[TEXT]{color} or {color:name}\[TEXT\]{color}
    # Note: brackets may be escaped in markdown from Confluence
    status_pattern = re.compile(r'\{color:(\w+)\}\\?\[([^\]\\]+)\\?\]\{color\}')

    def status_replacer(match):
        color = match.group(1).lower()
        text = match.group(2)
        return f'<span class="status status-{color}">{text}</span>'

    result = status_pattern.sub(status_replacer, result)

    # Convert text emoticons and Confluence-style codes to emoji
    result = replace_emoticons(result)

    # Convert single-tilde strikethrough to HTML <del>
    result = re.sub(r'~([^~]+)~', r'<del>\1</del>', result)

    # Remove anchor macros from output
    result = re.sub(r'\{anchor:[^}]+\}', '', result)

    return result


def replace_emoticons(text: str) -> str:
    """Replace common text emoticons with emoji."""
    # Confluence-style :name: codes
    name_map = {
        "smile": "😊",
        "sad": "🙁",
        "laugh": "😄",
        "wink": "😉",
        "thumbs-up": "👍",
        "thumbs_down": "👎",
        "information": "ℹ️",
        "warning": "⚠️",
        "question": "❓",
        "check": "✅",
        "cross": "❌",
        "star": "⭐",
    }

    def name_replacer(match):
        name = match.group(1).lower()
        return name_map.get(name, match.group(0))

    text = re.sub(r':([a-zA-Z0-9_-]+):', name_replacer, text)

    # Simple emoticons and Confluence shorthand
    simple_map = {
        ":)": "😊",
        ":-)": "😊",
        ":(": "🙁",
        ":-(": "🙁",
        ":D": "😄",
        ":-D": "😄",
        ";)": "😉",
        ";-)": "😉",
    }

    for key, emoji in simple_map.items():
        text = re.sub(rf'(^|[\s]){re.escape(key)}(?=$|[\s\.,;:!\?])', rf'\1{emoji}', text)

    shorthand_map = {
        "y": "👍",
        "n": "👎",
        "i": "ℹ️",
        "!": "⚠️",
        "?": "❓",
        "/": "✅",
        "x": "❌",
        "*": "⭐",
    }

    def shorthand_replacer(match):
        prefix = match.group(1)
        code = match.group(2).lower()
        return prefix + shorthand_map.get(code, match.group(0))

    text = re.sub(r'(^|[\s])\(\s*\\?([ynix\*/\?!])\s*\)(?=$|[\s\.,;:!\?])', shorthand_replacer, text, flags=re.IGNORECASE)

    return text


class MarkdownToWordConverter:
    """Converts markdown text to Word document elements."""

    def __init__(
        self,
        template: DocxTemplate,
        images: dict | None = None,
        macro_children: list | None = None,
        content_by_label: list | None = None,
    ):
        """Initialize the converter.

        Args:
            template: DocxTemplate instance for creating subdocuments
            images: Dict of embedded images keyed by filename.
                    Each value is {"data": base64_string, "mimeType": "image/png"}
        """
        self.template = template
        self.images = images or {}
        self.macro_children = macro_children or []
        self.content_by_label = content_by_label or []
        self.heading_numbered = self._detect_numbered_headings()
        self.headings: list[tuple[int, str]] = []
        self.md = markdown.Markdown(
            extensions=[
                TableExtension(),
                FencedCodeExtension(),
                "nl2br",
            ]
        )

    def convert(self, md_text: str) -> str:
        """Convert markdown to subdocument for template insertion.

        For the vertical slice, we return the content as a subdocument
        that can be inserted into the template.
        """
        # Extract headings for TOC generation before preprocessing
        self.headings = self._extract_headings(md_text)

        # Pre-process custom syntax (panels, status badges)
        preprocessed = preprocess_markdown(md_text)

        # Parse markdown to HTML
        html = self.md.convert(preprocessed)
        self.md.reset()

        # Create subdocument
        subdoc = self.template.new_subdoc()

        # Parse HTML and build Word elements
        soup = BeautifulSoup(html, "html.parser")
        self._process_elements(soup, subdoc)

        return subdoc

    def _extract_headings(self, md_text: str) -> list[tuple[int, str]]:
        """Extract markdown headings for TOC generation."""
        headings: list[tuple[int, str]] = []
        def unescape(text: str) -> str:
            return re.sub(r'\\([\\`*_{}\[\]()+\-\.!])', r'\1', text)

        for line in md_text.splitlines():
            if line.startswith("#"):
                m = re.match(r'^(#{1,6})\s+(.*)$', line)
                if not m:
                    continue
                level = len(m.group(1))
                text = unescape(m.group(2).strip())
                if not text:
                    continue
                headings.append((level, text))
        return headings

    def _detect_numbered_headings(self) -> dict[int, bool]:
        """Detect whether Heading 1-6 styles are numbered in the template."""
        numbered: dict[int, bool] = {}
        doc = self.template.get_docx()
        for level in range(1, 7):
            style = self._find_heading_style(doc, level)
            numbered[level] = self._style_has_numbering(style) if style else False
        return numbered

    def _find_heading_style(self, doc, level: int):
        target = f"heading {level}"
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.lower() == target:
                return style
        return None

    def _style_has_numbering(self, style) -> bool:
        current = style
        seen = set()
        while current is not None and id(current) not in seen:
            seen.add(id(current))
            ppr = current.element.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:numPr")) is not None:
                return True
            current = current.base_style
        return False

    def _strip_heading_prefix(self, text: str) -> tuple[str, bool]:
        """Strip leading numeric prefix like '6. ' or '1.2. '."""
        pattern = re.compile(r'^\s*\d+(?:\.\d+)*\.\s+')
        stripped = pattern.sub("", text, count=1)
        return stripped, stripped != text

    def _process_elements(self, soup: BeautifulSoup, subdoc) -> None:
        """Process HTML elements and add to subdocument."""
        for element in soup.children:
            if isinstance(element, Tag):
                self._process_tag(element, subdoc)
            elif isinstance(element, str) and element.strip():
                # Plain text
                p = subdoc.add_paragraph(element.strip())

    def _process_tag(self, tag: Tag, subdoc, paragraph=None) -> None:
        """Process a single HTML tag."""
        tag_name = tag.name.lower()

        if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag_name[1])
            p = subdoc.add_paragraph(style=f"Heading {level}")
            strip_number = self.heading_numbered.get(level, False)
            self._add_inline_content(tag, p, strip_leading_number=strip_number)

        elif tag_name == "p":
            p = subdoc.add_paragraph()
            self._add_inline_content(tag, p)

        elif tag_name == "ul":
            self._process_list(tag, subdoc, ordered=False)

        elif tag_name == "ol":
            self._process_list(tag, subdoc, ordered=True)

        elif tag_name == "pre":
            # Code block with better styling
            code = tag.find("code")
            code_text = code.get_text() if code else tag.get_text()
            p = subdoc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            # Light gray background and border
            self._add_code_block_style(p)

        elif tag_name == "img":
            # Block-level image
            src = tag.get("src", "")
            alt = tag.get("alt", "")
            p = subdoc.add_paragraph()
            self._add_image(p, src, alt)

        elif tag_name == "blockquote":
            for child in tag.children:
                if isinstance(child, Tag) and child.name == "p":
                    p = subdoc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    self._add_inline_content(child, p)
                    # Add italic styling for quotes
                    for run in p.runs:
                        run.italic = True

        elif tag_name == "table":
            self._process_table(tag, subdoc)

        elif tag_name == "hr":
            # Horizontal rule - add empty paragraph with bottom border
            p = subdoc.add_paragraph()

        elif tag_name == "div":
            # Check for panel macro
            classes = tag.get("class", [])

            if "toc-macro" in classes:
                self._insert_toc_macro_output(subdoc)
                return

            if "expand" in classes:
                self._process_expand(tag, subdoc)
                return

            if "children" in classes:
                self._process_children_macro(subdoc)
                return

            if "content-by-label" in classes:
                self._process_content_by_label_macro(tag, subdoc)
                return

            panel_type = None
            for cls in classes:
                if cls.startswith("panel-") and cls != "panel-title":
                    panel_type = cls.replace("panel-", "")
                    break

            if panel_type and panel_type in PANEL_STYLES:
                self._process_panel(tag, subdoc, panel_type)
            else:
                # Regular div - process children
                for child in tag.children:
                    if isinstance(child, Tag):
                        self._process_tag(child, subdoc)
                    elif isinstance(child, str) and child.strip():
                        p = subdoc.add_paragraph(child.strip())

        elif tag_name in ("section", "article"):
            # Container elements - process children
            for child in tag.children:
                if isinstance(child, Tag):
                    self._process_tag(child, subdoc)

        else:
            # Unknown block element - try to process as paragraph
            p = subdoc.add_paragraph()
            self._add_inline_content(tag, p)

    def _add_inline_content(
        self,
        tag: Tag,
        paragraph,
        strip_leading_number: bool = False,
        strip_state: dict | None = None,
    ) -> None:
        """Add inline content (text, bold, italic, links, code) to a paragraph."""
        if strip_state is None:
            strip_state = {"done": False}

        if tag.name and tag.name.lower() == "a":
            href = tag.get("href", "")
            text = tag.get_text()
            if strip_leading_number and not strip_state["done"]:
                text, stripped = self._strip_heading_prefix(text)
                if stripped:
                    strip_state["done"] = True
            self._add_hyperlink(paragraph, href, text)
            return

        for child in tag.children:
            if isinstance(child, str):
                text = child
                if strip_leading_number and not strip_state["done"]:
                    text, stripped = self._strip_heading_prefix(text)
                    if stripped:
                        strip_state["done"] = True
                if text.strip():
                    paragraph.add_run(text)
            elif isinstance(child, Tag):
                child_name = child.name.lower()

                if child_name in ("strong", "b"):
                    text = child.get_text()
                    if strip_leading_number and not strip_state["done"]:
                        text, stripped = self._strip_heading_prefix(text)
                        if stripped:
                            strip_state["done"] = True
                    run = paragraph.add_run(text)
                    run.bold = True

                elif child_name in ("em", "i"):
                    text = child.get_text()
                    if strip_leading_number and not strip_state["done"]:
                        text, stripped = self._strip_heading_prefix(text)
                        if stripped:
                            strip_state["done"] = True
                    run = paragraph.add_run(text)
                    run.italic = True

                elif child_name == "code":
                    text = child.get_text()
                    if strip_leading_number and not strip_state["done"]:
                        text, stripped = self._strip_heading_prefix(text)
                        if stripped:
                            strip_state["done"] = True
                    run = paragraph.add_run(text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)

                elif child_name == "a":
                    # Proper Word hyperlink
                    href = child.get("href", "")
                    text = child.get_text()
                    if strip_leading_number and not strip_state["done"]:
                        text, stripped = self._strip_heading_prefix(text)
                        if stripped:
                            strip_state["done"] = True
                    self._add_hyperlink(paragraph, href, text)

                elif child_name == "br":
                    paragraph.add_run("\n")

                elif child_name == "img":
                    # Handle inline image
                    src = child.get("src", "")
                    alt = child.get("alt", "")
                    self._add_image(paragraph, src, alt)

                elif child_name == "span":
                    # Check for status badge
                    classes = child.get("class", [])
                    if "status" in classes:
                        self._add_status_badge(child, paragraph)
                    else:
                        # Recursively process span content
                        self._add_inline_content(
                            child,
                            paragraph,
                            strip_leading_number=strip_leading_number,
                            strip_state=strip_state,
                        )

                elif child_name == "del":
                    text = child.get_text()
                    if strip_leading_number and not strip_state["done"]:
                        text, stripped = self._strip_heading_prefix(text)
                        if stripped:
                            strip_state["done"] = True
                    run = paragraph.add_run(text)
                    run.font.strike = True

                else:
                    # Recursively process other inline elements
                    self._add_inline_content(
                        child,
                        paragraph,
                        strip_leading_number=strip_leading_number,
                        strip_state=strip_state,
                    )

    def _process_list(self, tag: Tag, subdoc, ordered: bool = False, level: int = 0) -> None:
        """Process ordered or unordered list."""
        for i, li in enumerate(tag.find_all("li", recursive=False)):
            p = subdoc.add_paragraph()

            # Detect task list marker at start
            task_match = re.match(r'^\s*\[([ xX])\]\s*', li.get_text(strip=True))
            checkbox = None
            if task_match:
                checkbox = "\u2610 " if task_match.group(1) == " " else "\u2611 "

            # Add bullet or number (skip bullet if task list to avoid double markers)
            if checkbox:
                prefix = checkbox
            elif ordered:
                prefix = f"{i + 1}. "
            else:
                prefix = "\u2022 "  # Bullet character

            # Indent based on nesting level
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))

            # Add prefix
            p.add_run(prefix)

            # Process content
            for child in li.children:
                if isinstance(child, str):
                    if child.strip():
                        text = child.strip()
                        if checkbox:
                            text = re.sub(r'^\[[ xX]\]\s*', '', text)
                        if text:
                            p.add_run(text)
                elif isinstance(child, Tag):
                    if child.name in ("ul", "ol"):
                        # Nested list
                        self._process_list(
                            child, subdoc,
                            ordered=(child.name == "ol"),
                            level=level + 1
                        )
                    elif child.name == "a":
                        href = child.get("href", "")
                        text = child.get_text()
                        self._add_hyperlink(p, href, text)
                    else:
                        self._add_inline_content(child, p)

    def _process_table(self, tag: Tag, subdoc) -> None:
        """Process HTML table to Word table."""
        rows = tag.find_all("tr")
        if not rows:
            return

        # Count columns from first row
        first_row = rows[0]
        cols = len(first_row.find_all(["th", "td"]))

        if cols == 0:
            return

        # Create table
        table = subdoc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"

        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(["th", "td"])
            for col_idx, cell in enumerate(cells):
                if col_idx < cols:
                    word_cell = table.rows[row_idx].cells[col_idx]
                    word_cell.text = cell.get_text().strip()

                    # Bold for header cells
                    if cell.name == "th":
                        for paragraph in word_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _process_panel(self, tag: Tag, subdoc, panel_type: str) -> None:
        """Process a panel macro (info, warning, note, tip) as a styled box."""
        style = PANEL_STYLES.get(panel_type, PANEL_STYLES["info"])

        # Find title if present
        title_div = tag.find("div", class_="panel-title")
        title = title_div.get_text().strip() if title_div else None

        # Create a table with 1 cell to simulate a box with border/background
        table = subdoc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Style the cell with background and border
        self._style_panel_cell(cell, style)

        # Add title if present
        if title:
            title_p = cell.paragraphs[0]
            title_run = title_p.add_run(f"{style['icon']} {title}")
            title_run.bold = True
            title_run.font.size = Pt(11)
        else:
            # Add icon only
            first_p = cell.paragraphs[0]
            icon_run = first_p.add_run(f"{style['icon']} ")

        # Process content (skip the title div)
        for child in tag.children:
            if isinstance(child, Tag):
                if "panel-title" in child.get("class", []):
                    continue  # Skip title, already handled

                # For paragraphs, add to existing or new paragraph in cell
                if child.name == "p":
                    # If first paragraph is empty (no title), use it
                    if not title and len(cell.paragraphs) == 1 and not cell.paragraphs[0].text.strip():
                        p = cell.paragraphs[0]
                        p.add_run(f"{style['icon']} ")
                    else:
                        p = cell.add_paragraph()
                    self._add_inline_content(child, p)
                elif child.name in ("ul", "ol"):
                    self._process_list_in_cell(child, cell, ordered=(child.name == "ol"))
                else:
                    # For other elements, add as new paragraph
                    p = cell.add_paragraph()
                    self._add_inline_content(child, p)
            elif isinstance(child, str):
                text = child.strip()
                if not text:
                    continue
                # Plain text inside panel (no wrapping <p>)
                if not title and len(cell.paragraphs) == 1 and not cell.paragraphs[0].text.strip():
                    p = cell.paragraphs[0]
                    p.add_run(f"{style['icon']} ")
                else:
                    p = cell.add_paragraph()
                p.add_run(text)

        # Add spacing after panel
        subdoc.add_paragraph()

    def _style_panel_cell(self, cell, style: dict) -> None:
        """Apply panel styling to a table cell."""
        # Set cell background color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Background shading
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), style["bg"])
        tcPr.append(shading)

        # Cell borders
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "12")  # 1.5pt border
            border.set(qn("w:color"), style["border"])
            tcBorders.append(border)
        tcPr.append(tcBorders)

        # Cell margins/padding
        tcMar = OxmlElement("w:tcMar")
        for margin_name in ["top", "left", "bottom", "right"]:
            margin = OxmlElement(f"w:{margin_name}")
            margin.set(qn("w:w"), "144")  # ~0.1 inch
            margin.set(qn("w:type"), "dxa")
            tcMar.append(margin)
        tcPr.append(tcMar)

    def _add_shading(self, paragraph, color: str) -> None:
        """Add background shading to a paragraph."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        paragraph._p.get_or_add_pPr().append(shading)

    def _add_code_block_style(self, paragraph) -> None:
        """Add code block styling with background and border."""
        pPr = paragraph._p.get_or_add_pPr()

        # Add light gray background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), "F5F5F5")
        pPr.append(shading)

        # Add border around the paragraph
        pBdr = OxmlElement("w:pBdr")

        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")  # 0.5pt border
            border.set(qn("w:space"), "1")
            border.set(qn("w:color"), "CCCCCC")
            pBdr.append(border)

        pPr.append(pBdr)

        # Add padding via indentation
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "284")  # ~0.2 inch
        ind.set(qn("w:right"), "284")
        pPr.append(ind)

    def _add_hyperlink(self, paragraph, url: str, text: str) -> None:
        """Add a proper clickable hyperlink to a paragraph."""
        # Get the document part for relationship
        part = paragraph.part

        # Create relationship for the hyperlink
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create run with styled text
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Apply hyperlink style for consistent Word rendering
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        # Blue color
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)

        # Underline
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        new_run.append(rPr)

        # Add text
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_image(self, paragraph, src: str, alt: str = "") -> None:
        """Add an image to a paragraph.

        Args:
            paragraph: The paragraph to add the image to
            src: Image source - can be filename (for embedded) or URL
            alt: Alt text for the image
        """
        # Extract filename from src (may be URL path or just filename)
        if "/" in src:
            filename = src.split("/")[-1]
            # Remove query parameters
            if "?" in filename:
                filename = filename.split("?")[0]
        else:
            filename = src

        # Check if we have this image embedded
        if filename in self.images:
            image_data = self.images[filename]
            try:
                # Decode base64 image
                img_bytes = base64.b64decode(image_data["data"])
                img_stream = io.BytesIO(img_bytes)

                # Add image to document
                # Default max width of 5 inches, maintaining aspect ratio
                run = paragraph.add_run()
                run.add_picture(img_stream, width=Inches(5))

            except Exception as e:
                # If image fails, add placeholder text
                run = paragraph.add_run(f"[Image: {alt or filename}]")
                run.italic = True
        else:
            # No embedded image - add placeholder with alt text
            run = paragraph.add_run(f"[Image: {alt or filename}]")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_status_badge(self, tag: Tag, paragraph) -> None:
        """Add a colored status badge with background."""
        text = tag.get_text().strip()
        color_name = None

        # Extract color from class (status-green, status-red, etc.)
        for cls in tag.get("class", []):
            if cls.startswith("status-") and cls != "status":
                color_name = cls.replace("status-", "")
                break

        # Get color from map, default to gray
        color = STATUS_COLORS.get(color_name, STATUS_COLORS["gray"])

        # Add space before badge
        paragraph.add_run(" ")

        # Create the badge text with styling
        run = paragraph.add_run(f" {text} ")
        run.bold = True
        if color_name in ("gray", "grey", "yellow"):
            run.font.color.rgb = RGBColor(0, 0, 0)  # Dark text for light badges
        else:
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        run.font.size = Pt(9)

        # Add background highlight using shading on the run
        # Note: Word doesn't support per-run background easily,
        # so we use a highlight color approximation
        self._add_run_shading(run, color_name)

        # Add space after badge
        paragraph.add_run(" ")

    def _add_run_shading(self, run, color_name: str) -> None:
        """Add background shading to a run (approximated with highlight)."""
        # Map our colors to Word highlight colors
        highlight_map = {
            "green": "green",
            "yellow": "yellow",
            "red": "red",
            "blue": "blue",
            "grey": "lightGray",
            "gray": "lightGray",
            "purple": "darkMagenta",
        }

        highlight = highlight_map.get(color_name, "darkGray")

        # Apply highlight
        rPr = run._r.get_or_add_rPr()
        highlight_elem = OxmlElement("w:highlight")
        highlight_elem.set(qn("w:val"), highlight)
        rPr.append(highlight_elem)

    def _insert_toc_macro_output(self, subdoc) -> None:
        """Insert a Word-native TOC field wrapped in SDT container.

        The SDT (Structured Document Tag) wrapper allows the TOC to be marked
        as dirty so Word prompts the user to update fields on open.
        """
        # Create SDT container structure
        sdt = OxmlElement("w:sdt")

        # SDT properties - identifies this as a Table of Contents
        sdt_pr = OxmlElement("w:sdtPr")
        doc_part_obj = OxmlElement("w:docPartObj")
        doc_part_gallery = OxmlElement("w:docPartGallery")
        doc_part_gallery.set(qn("w:val"), "Table of Contents")
        doc_part_unique = OxmlElement("w:docPartUnique")
        doc_part_obj.append(doc_part_gallery)
        doc_part_obj.append(doc_part_unique)
        sdt_pr.append(doc_part_obj)
        sdt.append(sdt_pr)

        # SDT content - contains the actual TOC field
        sdt_content = OxmlElement("w:sdtContent")

        # Create paragraph for TOC field
        p_elem = OxmlElement("w:p")

        # Field begin run
        begin_run = OxmlElement("w:r")
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        begin_run.append(fld_begin)
        p_elem.append(begin_run)

        # Field instruction run: TOC with outline levels 1-3, hyperlinks
        instr_run = OxmlElement("w:r")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        instr_run.append(instr_text)
        p_elem.append(instr_run)

        # Field separate run
        sep_run = OxmlElement("w:r")
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        sep_run.append(fld_sep)
        p_elem.append(sep_run)

        # Placeholder text (will be replaced when Word updates the field)
        placeholder_run = OxmlElement("w:r")
        placeholder_text = OxmlElement("w:t")
        placeholder_text.text = "Table of Contents - Update to populate"
        placeholder_run.append(placeholder_text)
        p_elem.append(placeholder_run)

        # Field end run
        end_run = OxmlElement("w:r")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        end_run.append(fld_end)
        p_elem.append(end_run)

        # Assemble: paragraph into content, content into SDT
        sdt_content.append(p_elem)
        sdt.append(sdt_content)

        # Add a paragraph to subdoc and replace its XML with our SDT
        p = subdoc.add_paragraph()
        p._p.getparent().replace(p._p, sdt)

    def _process_expand(self, tag: Tag, subdoc) -> None:
        """Render expand macro as a styled box with title and content."""
        title = tag.get("data-title", "Click to expand")
        style = {"bg": "F5F5F5", "border": "7A869A", "icon": "▸"}

        table = subdoc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        self._style_panel_cell(cell, style)

        title_p = cell.paragraphs[0]
        title_run = title_p.add_run(f"{style['icon']} {title}")
        title_run.bold = True
        title_run.font.size = Pt(11)

        for child in tag.children:
            if isinstance(child, Tag):
                if child.name == "p":
                    p = cell.add_paragraph()
                    self._add_inline_content(child, p)
                elif child.name in ("ul", "ol"):
                    self._process_list_in_cell(child, cell, ordered=(child.name == "ol"))
                else:
                    p = cell.add_paragraph()
                    self._add_inline_content(child, p)
            elif isinstance(child, str) and child.strip():
                p = cell.add_paragraph()
                p.add_run(child.strip())

        subdoc.add_paragraph()

    def _process_children_macro(self, subdoc) -> None:
        """Render children macro as a bullet list."""
        if not self.macro_children:
            return

        for child in self.macro_children:
            title = child.get("title") or ""
            url = child.get("pageUrl") or ""
            p = subdoc.add_paragraph()
            p.add_run("\u2022 ")
            if url:
                self._add_hyperlink(p, url, title or url)
            else:
                p.add_run(title)

    def _process_content_by_label_macro(self, tag: Tag, subdoc) -> None:
        """Render content-by-label macro as a bullet list."""
        labels = tag.get("data-labels", "")
        spaces = tag.get("data-spaces", "")

        def normalize_list(value: str) -> list[str]:
            return [v.strip() for v in value.split(",") if v.strip()]

        labels_list = normalize_list(labels)
        spaces_list = normalize_list(spaces)

        match = None
        for entry in self.content_by_label:
            entry_labels = normalize_list(entry.get("labels", ""))
            entry_spaces = normalize_list(entry.get("spaces", ""))
            if set(entry_labels) == set(labels_list) and set(entry_spaces) == set(spaces_list):
                match = entry
                break

        items = match.get("items", []) if match else []
        if not items:
            return

        for item in items:
            title = item.get("title") or ""
            url = item.get("pageUrl") or ""
            p = subdoc.add_paragraph()
            p.add_run("\u2022 ")
            if url:
                self._add_hyperlink(p, url, title or url)
            else:
                p.add_run(title)

    def _process_list_in_cell(self, tag: Tag, cell, ordered: bool = False, level: int = 0) -> None:
        """Process list elements inside a table cell."""
        for i, li in enumerate(tag.find_all("li", recursive=False)):
            p = cell.add_paragraph()

            task_match = re.match(r'^\s*\[([ xX])\]\s*', li.get_text(strip=True))
            checkbox = None
            if task_match:
                checkbox = "\u2610 " if task_match.group(1) == " " else "\u2611 "

            if checkbox:
                prefix = checkbox
            elif ordered:
                prefix = f"{i + 1}. "
            else:
                prefix = "\u2022 "

            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            p.add_run(prefix)

            for child in li.children:
                if isinstance(child, str):
                    if child.strip():
                        text = child.strip()
                        if checkbox:
                            text = re.sub(r'^\[[ xX]\]\s*', '', text)
                        if text:
                            p.add_run(text)
                elif isinstance(child, Tag):
                    if child.name in ("ul", "ol"):
                        self._process_list_in_cell(
                            child,
                            cell,
                            ordered=(child.name == "ol"),
                            level=level + 1,
                        )
                    elif child.name == "a":
                        href = child.get("href", "")
                        text = child.get_text()
                        self._add_hyperlink(p, href, text)
                    else:
                        self._add_inline_content(child, p)
