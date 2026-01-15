"""Create a simple test template for testing the export functionality."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_test_template(output_path: Path) -> None:
    """Create a basic test template with Jinja2 placeholders.

    Note: docxtpl expects {{ variable }} syntax directly in the text.
    The placeholders must be written as literal text.
    """
    doc = Document()

    # Title - using literal Jinja2 syntax
    title = doc.add_heading("", level=0)
    title.add_run("{{ title }}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata section
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Author: ").bold = True
    meta.add_run("{{ author }}")
    meta.add_run(" | ")
    meta.add_run("Modified: ").bold = True
    meta.add_run("{{ modified | date('YYYY-MM-DD') }}")

    doc.add_paragraph()
    space_info = doc.add_paragraph()
    space_info.add_run("Space: ").bold = True
    space_info.add_run("{{ spaceName }} ({{ spaceKey }})")

    # Horizontal line (as a paragraph with bottom border)
    doc.add_paragraph("─" * 60)

    # Content section
    doc.add_heading("Content", level=1)
    # For subdocuments, we use special RichText marker
    content_para = doc.add_paragraph()
    content_para.add_run("{{r content }}")

    # Footer
    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph()
    footer.add_run("Exported: ").bold = True
    footer.add_run("{{ exportDate | date('YYYY-MM-DD HH:mm') }}")

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Created template: {output_path}")


if __name__ == "__main__":
    template_dir = Path(__file__).parent / "fixtures"
    create_test_template(template_dir / "basic-template.docx")
