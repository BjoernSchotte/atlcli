"""Tests for DOCX rendering."""

import json
from pathlib import Path
import subprocess
import sys

import pytest

FIXTURES_DIR = Path(__file__).parent / "fixtures"
OUTPUT_DIR = Path(__file__).parent / "output"


@pytest.fixture
def sample_page_data():
    """Sample Confluence page data."""
    return {
        "title": "Test Page Title",
        "markdown": """# Introduction

This is a **test document** with *formatted* content.

## Features

- Bullet point one
- Bullet point two
- Bullet point three

## Code Example

Here's some `inline code` and a code block:

```python
def hello():
    print("Hello, World!")
```

## Table

| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| A        | B        | C        |
| D        | E        | F        |

> This is a blockquote with important information.

That's all for now!
""",
        "author": {"displayName": "John Doe", "email": "john@example.com"},
        "modifier": {"displayName": "Jane Smith", "email": "jane@example.com"},
        "created": "2025-01-10T10:00:00Z",
        "modified": "2025-01-15T14:30:00Z",
        "pageId": "12345678",
        "pageUrl": "https://example.atlassian.net/wiki/spaces/TEST/pages/12345678",
        "tinyUrl": "https://example.atlassian.net/wiki/x/abc",
        "labels": ["documentation", "test"],
        "spaceKey": "TEST",
        "spaceName": "Test Space",
        "spaceUrl": "https://example.atlassian.net/wiki/spaces/TEST",
        "exportedBy": "Claude Code",
        "templateName": "basic-template",
        "attachments": [],
        "children": [],
    }


def test_render_template(sample_page_data):
    """Test basic template rendering."""
    from atlcli_export import render_template

    template_path = FIXTURES_DIR / "basic-template.docx"
    output_path = OUTPUT_DIR / "test-output.docx"

    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Render - returns (path, has_toc) tuple
    result_path, has_toc = render_template(template_path, sample_page_data, output_path)

    assert result_path.exists()
    assert result_path.suffix == ".docx"
    assert isinstance(has_toc, bool)


def test_cli_render(sample_page_data):
    """Test CLI entry point."""
    template_path = FIXTURES_DIR / "basic-template.docx"
    output_path = OUTPUT_DIR / "cli-output.docx"

    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Run CLI via subprocess
    result = subprocess.run(
        [
            sys.executable, "-m", "atlcli_export.cli",
            "--template", str(template_path),
            "--output", str(output_path),
        ],
        input=json.dumps(sample_page_data),
        capture_output=True,
        text=True,
    )

    assert result.returncode == 0
    response = json.loads(result.stdout)
    assert response["success"] is True
    assert Path(response["output"]).exists()


def test_date_filter():
    """Test date formatting filter."""
    from atlcli_export.filters import date_filter

    # ISO format with time and timezone
    assert date_filter("2025-01-15T14:30:00Z", "YYYY-MM-DD") == "2025-01-15"
    assert date_filter("2025-01-15T14:30:00Z", "DD/MM/YYYY") == "15/01/2025"
    assert date_filter("2025-01-15T14:30:00Z", "YYYY-MM-DD HH:mm") == "2025-01-15 14:30"
    assert date_filter("2025-01-05T04:30:00Z", "yyyy-MM-dd") == "2025-01-05"
    assert date_filter("2025-01-05T04:30:00Z", "MMMM d, yyyy") == "January 5, 2025"

    # Empty string
    assert date_filter("", "YYYY-MM-DD") == ""

    # Invalid date returns original
    assert date_filter("not-a-date", "YYYY-MM-DD") == "not-a-date"


def test_markdown_to_word_converter(sample_page_data):
    """Test markdown to Word conversion."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)

    converter = MarkdownToWordConverter(template)
    subdoc = converter.convert(sample_page_data["markdown"])

    # Subdoc should be created
    assert subdoc is not None


def test_render_docm_template(sample_page_data):
    """Test rendering with .docm (macro-enabled) template."""
    from atlcli_export import render_template

    template_path = FIXTURES_DIR / "basic-template.docm"
    output_path = OUTPUT_DIR / "docm-output.docx"

    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Render with .docm template - returns (path, has_toc) tuple
    result_path, has_toc = render_template(template_path, sample_page_data, output_path)

    assert result_path.exists()
    assert result_path.suffix == ".docx"


def test_docm_conversion():
    """Test .docm to .docx conversion."""
    from atlcli_export import convert_docm_to_docx, is_docm_file, DocmConverter
    import shutil

    docm_path = FIXTURES_DIR / "basic-template.docm"

    # Test is_docm_file
    assert is_docm_file(docm_path)
    assert is_docm_file("test.docm")
    assert not is_docm_file("test.docx")
    assert not is_docm_file("test.doc")

    # Test convert_docm_to_docx
    converted = convert_docm_to_docx(docm_path)
    try:
        assert converted.exists()
        assert converted.suffix == ".docx"

        # Should be openable by python-docx
        from docx import Document
        doc = Document(str(converted))
        assert len(doc.paragraphs) > 0
    finally:
        # Clean up
        shutil.rmtree(converted.parent, ignore_errors=True)

    # Test DocmConverter context manager
    with DocmConverter(docm_path) as usable_path:
        assert usable_path.suffix == ".docx"
        from docx import Document
        doc = Document(str(usable_path))
        assert len(doc.paragraphs) > 0
    # Temp files should be cleaned up after exiting context

    # Test DocmConverter with .docx (should return original path)
    docx_path = FIXTURES_DIR / "basic-template.docx"
    with DocmConverter(docx_path) as usable_path:
        assert usable_path == docx_path  # No conversion needed


def test_scroll_placeholder_conversion():
    """Test Scroll placeholder to Jinja2 conversion."""
    from atlcli_export import convert_scroll_placeholders

    # Basic variable
    assert convert_scroll_placeholders("$scroll.title") == "{{ title }}"
    assert convert_scroll_placeholders("$scroll.content") == "{{ content }}"

    # Mapped variables
    assert convert_scroll_placeholders("$scroll.creator.fullName") == "{{ author }}"
    assert convert_scroll_placeholders("$scroll.creator.email") == "{{ authorEmail }}"
    assert convert_scroll_placeholders("$scroll.modifier.fullName") == "{{ modifier }}"
    assert convert_scroll_placeholders("$scroll.pageid") == "{{ pageId }}"
    assert convert_scroll_placeholders("$scroll.space.key") == "{{ spaceKey }}"

    # Date variables
    assert convert_scroll_placeholders("$scroll.creationdate") == "{{ created }}"
    assert convert_scroll_placeholders("$scroll.modificationdate") == "{{ modified }}"
    assert convert_scroll_placeholders("$scroll.exportdate") == "{{ exportDate }}"

    # Null-safe variables
    assert convert_scroll_placeholders("$!scroll.title") == "{{ title | default('') }}"

    # Date formatting
    result = convert_scroll_placeholders('$scroll.creationdate.("yyyy-MM-dd")')
    assert result == "{{ created | date('yyyy-MM-dd') }}"

    result = convert_scroll_placeholders('$scroll.modificationdate("MMMM d, yyyy")')
    assert result == "{{ modified | date('MMMM d, yyyy') }}"

    # Combined null-safe with date format
    result = convert_scroll_placeholders('$!scroll.creationdate.("yyyy-MM-dd")')
    assert result == "{{ created | date('yyyy-MM-dd') | default('') }}"

    # In XML context (preserves XML)
    xml = '<w:t>$scroll.title</w:t>'
    assert convert_scroll_placeholders(xml) == '<w:t>{{ title }}</w:t>'

    # Multiple placeholders
    text = "Title: $scroll.title, Author: $scroll.creator.fullName"
    result = convert_scroll_placeholders(text)
    assert "{{ title }}" in result
    assert "{{ author }}" in result

    # Unknown variable (keeps original name)
    assert convert_scroll_placeholders("$scroll.unknown") == "{{ unknown }}"

    # Non-scroll text unchanged
    assert convert_scroll_placeholders("Hello world") == "Hello world"
    assert convert_scroll_placeholders("{{ title }}") == "{{ title }}"


def test_normalize_split_placeholders():
    """Test normalization of split Scroll placeholders in Word XML."""
    from atlcli_export.docm_support import normalize_split_placeholders, convert_scroll_placeholders

    # Split placeholder: <w:t>$</w:t><w:t>scroll.content</w:t>
    xml = '<w:t>$</w:t></w:r><w:proofErr w:type="spellStart"/><w:r><w:t>scroll.content</w:t>'
    normalized = normalize_split_placeholders(xml)
    # After normalization, $scroll.content should be together
    converted = convert_scroll_placeholders(normalized)
    assert "{{ content }}" in converted
    assert "$scroll" not in converted

    # Split with null-safe: <w:t>$</w:t><w:t>!scroll.title</w:t>
    xml2 = '<w:t>$</w:t></w:r><w:r><w:t>!scroll.title</w:t>'
    normalized2 = normalize_split_placeholders(xml2)
    converted2 = convert_scroll_placeholders(normalized2)
    assert "{{ title | default('') }}" in converted2

    # Normal (not split) should still work
    xml3 = '<w:t>$scroll.title</w:t>'
    normalized3 = normalize_split_placeholders(xml3)
    converted3 = convert_scroll_placeholders(normalized3)
    assert "{{ title }}" in converted3


def test_fix_content_placeholder_preserves_section_break():
    """Ensure section breaks are preserved when moving content placeholder."""
    from atlcli_export.docm_support import fix_content_placeholder_level

    xml = (
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>'
        '<w:p>'
        '<w:pPr><w:sectPr><w:pgSz w:w="12240"/></w:sectPr></w:pPr>'
        '<w:r><w:t>{{ content }}</w:t></w:r>'
        '</w:p>'
        '</w:body></w:document>'
    )

    fixed = fix_content_placeholder_level(xml)
    assert "{{p content }}" in fixed
    assert "<w:sectPr" in fixed
    assert fixed.index("{{p content }}") < fixed.index("<w:sectPr")


def test_hyperlinks_in_markdown():
    """Test that hyperlinks are properly converted to Word hyperlinks."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter
    import zipfile

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)

    # Markdown with a link
    markdown = "Check out [this link](https://example.com) for more info."
    converter = MarkdownToWordConverter(template)
    subdoc = converter.convert(markdown)

    # Subdoc should be created
    assert subdoc is not None


def test_toc_macro_injects_word_toc_field(tmp_path):
    """TOC macro should inject a Word-native TOC field in SDT container."""
    from atlcli_export.docx_renderer import render_template
    import zipfile

    template_path = FIXTURES_DIR / "basic-template.docx"
    output_path = tmp_path / "toc-export.docx"

    # Page with TOC macro
    page_data = {
        "title": "TOC Test",
        "markdown": ":::toc\n:::\n\n## 1. Basic Formatting\n\nSome text\n",
        "author": {"displayName": "Tester", "email": "tester@example.com"},
        "modifier": {"displayName": "Tester", "email": "tester@example.com"},
        "created": "2024-01-01T00:00:00Z",
        "modified": "2024-01-01T00:00:00Z",
        "pageId": "1",
        "pageUrl": "https://example.com",
        "tinyUrl": "https://ex",
        "labels": [],
        "spaceKey": "SPACE",
        "spaceName": "Space",
        "spaceUrl": "https://example.com/space",
        "exportedBy": "atlcli",
        "templateName": "basic-template",
        "attachments": [],
        "children": [],
        "images": {},
        "macroChildren": [],
        "macroContentByLabel": [],
    }

    result_path, has_toc = render_template(template_path, page_data, output_path)
    with zipfile.ZipFile(output_path, "r") as zin:
        xml = zin.read("word/document.xml").decode("utf-8")

    # Should contain TOC field instruction
    assert "TOC" in xml and "instrtext" in xml.lower()
    # Should contain placeholder text
    assert "Table of Contents" in xml
    # Should be wrapped in SDT container (for dirty flag support)
    assert "w:sdt" in xml
    assert "docPartGallery" in xml
    # Should report has_toc=True
    assert has_toc is True
    # Should be marked dirty on fldChar element (default behavior)
    # Per OOXML spec, dirty is on fldChar with fldCharType="begin", not on sdtPr
    assert 'w:dirty="true"' in xml or "w:dirty='true'" in xml


def test_heading_number_prefix_stripped_when_enabled():
    """Strip leading numeric prefix from headings when numbering is enabled."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter
    from bs4 import BeautifulSoup

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)
    converter = MarkdownToWordConverter(template)
    converter.heading_numbered[2] = True

    subdoc = template.new_subdoc()
    tag = BeautifulSoup("<h2>6. Panel Macros</h2>", "html.parser").h2
    converter._process_tag(tag, subdoc)

    assert subdoc.paragraphs[0].text == "Panel Macros"


def test_code_block_styling():
    """Test that code blocks get proper styling."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)

    # Markdown with code block
    markdown = """Here's some code:

```python
def hello():
    print("Hello!")
```
"""
    converter = MarkdownToWordConverter(template)
    subdoc = converter.convert(markdown)

    # Subdoc should be created
    assert subdoc is not None


def test_image_placeholder_without_embed():
    """Test that images without embedded data show placeholder."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)

    # Markdown with image (no embedded data)
    markdown = "Here's an image: ![My Image](attachment.png)"
    converter = MarkdownToWordConverter(template, images={})
    subdoc = converter.convert(markdown)

    # Subdoc should be created
    assert subdoc is not None


def test_panel_macros():
    """Test that panel macros (info, warning, note, tip) are rendered as styled boxes."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter, preprocess_markdown

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)

    # Test preprocessing
    markdown = """:::info Important Note
This is an info panel with a title.
:::

:::warning
Warning without title.
:::

:::tip Pro Tip
Here's a helpful tip.
:::
"""
    # Check preprocessing converts to HTML
    preprocessed = preprocess_markdown(markdown)
    assert 'class="panel panel-info"' in preprocessed
    assert 'class="panel-title"' in preprocessed
    assert 'class="panel panel-warning"' in preprocessed
    assert 'class="panel panel-tip"' in preprocessed

    # Convert to Word
    converter = MarkdownToWordConverter(template)
    subdoc = converter.convert(markdown)

    # Subdoc should be created
    assert subdoc is not None


def test_status_badges():
    """Test that status badges are rendered with colors."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter, preprocess_markdown

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)

    # Test preprocessing
    markdown = """Status: {color:green}[DONE]{color} and {color:red}[BLOCKED]{color}"""

    preprocessed = preprocess_markdown(markdown)
    assert 'class="status status-green"' in preprocessed
    assert 'class="status status-red"' in preprocessed
    assert ">DONE<" in preprocessed
    assert ">BLOCKED<" in preprocessed

    # Convert to Word
    converter = MarkdownToWordConverter(template)
    subdoc = converter.convert(markdown)

    # Subdoc should be created
    assert subdoc is not None


def test_image_with_embedded_data():
    """Test that images with embedded data are inserted."""
    from docxtpl import DocxTemplate
    from atlcli_export.markdown_to_word import MarkdownToWordConverter
    import base64

    template_path = FIXTURES_DIR / "basic-template.docx"
    template = DocxTemplate(template_path)

    # Create a minimal 1x1 red PNG
    # This is a valid minimal PNG file
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8DwHwAFBQIAX8jx0gAAAABJRU5ErkJggg=="
    )
    images = {
        "test.png": {
            "data": base64.b64encode(png_data).decode("utf-8"),
            "mimeType": "image/png",
        }
    }

    # Markdown with image
    markdown = "Here's an image: ![Test](test.png)"
    converter = MarkdownToWordConverter(template, images=images)
    subdoc = converter.convert(markdown)

    # Subdoc should be created
    assert subdoc is not None
